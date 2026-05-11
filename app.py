# ═══════════════════════════════════════════════════════════════════
# Astra v1.0 — Vishnunath Nharekkat
# Data Analyst, BI, Junior Data Scientist, Analytics Engineer Resume Tailoring Engine
#
# v1.0 (May 2026):
#  - 8 JD archetypes covering Vishnu's real Dublin market:
#      data_analyst | bi_analyst | junior_data_scientist | analytics_engineer
#      graduate_programme_tech | graduate_programme_business
#      qa_automation | unknown_stretch
#  - Stamp 1G surfaced as a CSEP-eligibility differentiator (€40,904 floor as of Mar 2026)
#  - Permissive banned-skills: in-progress upskilling certs (PL-300, T-SQL Advanced,
#    Tableau Public, AZ-900) only render with explicit '(In Progress)' marker
#  - Lighter cover letter: single 4-paragraph prompt, no war-story branching
#  - Projects section is separate from Professional Experience (honesty: a project
#    is not a job)
#
# Design goals (matching Praghya's app.py):
#  - Simple. No conditional gates that block CV generation.
#  - Truthful. Tailor to JD using ONLY skills/experience he actually has.
#  - Structure-preserving. Every base-resume section is always present.
#  - All his real base skills are kept; JD-relevant skills are added on top.
#  - Education / Certifications / Additional Info: untouched, always rendered.
#  - ATS-friendly output every run.
# ═══════════════════════════════════════════════════════════════════

import streamlit as st
import json
import re
import io
import datetime
from typing import List, Optional
from pydantic import BaseModel, Field
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from xml.sax.saxutils import escape


# ═══════════════════════════════════════════════════════════════════
# 1. CONFIG
# ═══════════════════════════════════════════════════════════════════

PAGE_TITLE = "Astra — Vishnunath Nharekkat"

# Generation model.
# As of May 2026, gemini-3-flash-preview is the recommended free-tier model
# in the Gemini 3 family. It handles structured-JSON output reliably.
#
# Migration notes — if you see errors, swap MODEL string to one of:
#   - "gemini-3.1-flash-lite-preview"  (cheaper, slightly weaker reasoning)
#   - "gemini-2.5-flash"               (stable until June 17, 2026)
#   - "gemini-2.5-pro"                 (best quality, lower rate limits, deprecates June 17, 2026)
GENERATION_MODEL = "gemini-3-flash-preview"

try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
except Exception:
    GOOGLE_API_KEY = ""


# ═══════════════════════════════════════════════════════════════════
# 2. FIXED CANDIDATE FACTS
# These NEVER change between runs, no matter what the JD says.
# ═══════════════════════════════════════════════════════════════════

CANDIDATE_NAME = "Vishnunath Nharekkat"
CANDIDATE_TAGLINE = "Data & Analytics Professional"
CANDIDATE_CONTACT = "Dublin, Ireland | +353 89 958 7478 | nvishnunath2706@gmail.com"

# Base skill set — real and verified. These are ALWAYS in the output.
# Each category is a list of skills. The tailoring step may APPEND
# JD-relevant skills (that he genuinely has), but it cannot remove any of these.
BASE_SKILLS = {
    "Languages and Querying": [
        "Python (pandas, NumPy, scikit-learn, matplotlib, seaborn)",
        "SQL", "R (basic)",
    ],
    "Business Intelligence and Visualisation": [
        "Power BI", "DAX", "Power Query (M)",
        "Microsoft Excel (PivotTables, VLOOKUP, XLOOKUP, INDEX/MATCH)",
    ],
    "Machine Learning and Statistics": [
        "Regression", "Classification", "Clustering",
        "Feature Engineering", "Hyperparameter Tuning", "Model Evaluation",
        "Hypothesis Testing", "A/B Testing",
        "NLP", "TF-IDF", "LSTM",
    ],
    "Data Processes": [
        "Data Cleansing", "Data Validation", "Data Quality Assessment",
        "ETL Fundamentals", "KPI Reporting", "Dashboard Development",
        "Star Schema", "Data Modelling",
    ],
    "Tools and Productivity": [
        "Git", "GitHub", "Jupyter", "JIRA",
        "Microsoft Office Suite (Word, PowerPoint)", "Confluence",
    ],
    "Geospatial and Engineering Tools": [
        "QGIS", "MapInfo", "AutoCAD",
    ],
    "Methods and Soft Skills": [
        "Stakeholder Management", "Cross-functional Collaboration",
        "Requirements Gathering", "Data Storytelling",
        "Attention to Detail", "Problem Solving",
    ],
}

# One real job. Tailoring rewrites the BULLETS, but never adds/removes the role
# or changes company/dates/title.
BASE_EXPERIENCE = [
    {
        "role_title": "Telecom Engineer",
        "company": "IMMCO",
        "location": "Kochi, India",
        "dates": "June 2021 - September 2023",
        "responsibilities": [
            "Performed geospatial data analysis and validation across large fibre network datasets in QGIS and MapInfo to support FTTP and FTTH design across SDU and MDU premises, applying spatial querying and routing logic at design stage.",
            "Ran Quality Assurance and First Pass validation on fibre routing, splitter placement, and splicing logic; logged defects in the design data and fed corrections back into the build workflow before construction handover.",
            "Built KPI tracking workbooks in Microsoft Excel that compared design deliverables (IDD, UDD, CP, ADD) against programme milestones and budget forecasts, and produced weekly status reports for project managers.",
            "Translated GPON architecture requirements into structured design specifications and reconciled data inconsistencies across AutoCAD, MapInfo, and QGIS outputs to confirm a single source of truth for each route.",
            "Acted as Subject Matter Expert and informal team lead for 4 to 6 design engineers; ran daily stand-ups, standardised validation checklists, and mentored junior engineers on data quality procedures.",
            "Coordinated with project and build managers across N2P and RCP fibre rollout programmes to track issue resolution and own cross-functional handovers between design, planning, and field teams.",
        ],
        "achievements": [
            "Reduced design errors and downstream rework by approximately 20 percent through structured QA and First Pass validation across N2P and RCP fibre programmes.",
            "Acted as Subject Matter Expert for OSP design validation across the team, owning the QA workflow and onboarding new engineers into validation procedures.",
            "Standardised data validation checklists and review templates that became the team default for OSP and UDD packages on the FTTP programme.",
        ],
    },
]

# Projects are rendered in their own section, distinct from Professional Experience.
# A project is not a job; keeping them separate preserves honesty.
BASE_PROJECTS = [
    {
        "title": "Comparative Machine Learning for Political Tweet Sentiment Analysis",
        "context": "MSc Thesis, NCI",
        "dates": "January 2025 - April 2025",
        "bullets": [
            "Built an end-to-end NLP pipeline in Python to classify sentiment in Indian election tweets (X / Twitter) on Narendra Modi, BJP, Rahul Gandhi, and INC; the pipeline covered scraping, preprocessing, tokenisation, lemmatisation, and feature engineering.",
            "Compared a TF-IDF and Linear Support Vector Machine baseline (86 percent accuracy) against a Long Short-Term Memory deep learning model (87 percent accuracy); evaluation used confusion matrices, F1 scores per class, and class-imbalance analysis.",
            "Worked in pandas, NumPy, scikit-learn, NLTK, and TensorFlow / Keras for modelling and evaluation; wrote up the full experimental design and results in a 20,000 word dissertation graded at 2:1.",
        ],
    },
    {
        "title": "End-to-End Power BI Sales and Operations Dashboard",
        "context": "Self-Directed Project",
        "dates": "2026",
        "bullets": [
            "Designed a star-schema data model in SQL on a public retail dataset to support sales, operations, and finance reporting, with fact and dimension tables for orders, customers, products, and time.",
            "Built a multi-page Power BI report with DAX measures for revenue, margin, customer cohort, and KPI variance against target; used Power Query (M) for data ingestion and cleansing.",
            "Added role-based filtering and drill-through pages so business users could move from headline KPIs to underlying transaction-level data without writing SQL.",
        ],
    },
]

# Education, certifications, additional info — NEVER tailored. Always rendered as-is.
BASE_EDUCATION = [
    {
        "degree": "Master of Science in Data Analytics",
        "institution": "National College of Ireland, Dublin",
        "dates": "January 2024 - April 2025",
        "grade": "Grade: 2:1",
        "extra": "Thesis: From Traditional to Advanced Machine Learning, a Comparative Study of Political Tweet Sentiment Analysis (TF-IDF + SVM vs LSTM).",
    },
    {
        "degree": "Bachelor of Science in Electronics",
        "institution": "V V College of Science and Technology, University of Calicut, India",
        "dates": "2017 - 2020",
        "grade": "Grade: First Class",
        "extra": "",
    },
]

BASE_CERTIFICATIONS = [
    "Power BI Data Analytics, Codebasics (2025)",
    "SQL (Basic), HackerRank (2025)",
    "Excel: Mother of Business Intelligence, Codebasics (2025)",
    "Python, Data Science, Machine Learning, AI and Power BI, Luminar Technolab (2025)",
    "Python 101 for Data Science, IBM Cognitive Class (2024)",
]

BASE_ADDITIONAL_INFO = [
    "Languages: English (Fluent), Malayalam (Native), Hindi (Conversational), Tamil (Conversational)",
    "Currently Learning: PL-300 (Microsoft Power BI Data Analyst), T-SQL Advanced (HackerRank), Tableau Public, AZ-900 (Azure Fundamentals)",
    "Work Authorisation: Stamp 1G (Third Level Graduate Programme), valid through April 2027. Eligible for Critical Skills Employment Permit. No immediate employment permit required.",
]


def build_base_resume_text() -> str:
    """Compose Vishnu's base resume as plain text from the constants above.
    This is what pre-fills the 'Base Resume' box in the UI so he can see
    and edit what's being sent to the model."""
    lines = []
    lines.append(CANDIDATE_NAME)
    lines.append(CANDIDATE_TAGLINE)
    lines.append(CANDIDATE_CONTACT)
    lines.append("")

    lines.append("Professional Profile")
    lines.append(
        "MSc Data Analytics graduate (2:1, National College of Ireland) with 2 years "
        "and 3 months of engineering experience in quantitative validation, geospatial "
        "data analysis, and KPI-led programme tracking. Hands-on with Python (pandas, "
        "NumPy, scikit-learn), SQL, Power BI, and advanced Excel for data cleansing, "
        "dashboard development, and statistical analysis. MSc thesis applied NLP and "
        "deep learning (TF-IDF with SVM, and LSTM) to political tweet sentiment "
        "classification, with the LSTM model reaching 87 percent accuracy. Open to "
        "Data Analyst, Business Intelligence, and Analytics Engineer roles across "
        "Ireland. Stamp 1G valid through April 2027 and eligible for a Critical "
        "Skills Employment Permit."
    )
    lines.append("")

    lines.append("Key Skills / Tools & Technologies")
    for cat, skill_list in BASE_SKILLS.items():
        lines.append(f"- {cat}: {', '.join(skill_list)}")
    lines.append("")

    lines.append("Professional Experience")
    for role in BASE_EXPERIENCE:
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        lines.append(header)
        for r in role["responsibilities"]:
            lines.append(f"- {r}")
        if role["achievements"]:
            lines.append("Achievements:")
            for a in role["achievements"]:
                lines.append(f"- {a}")
        lines.append("")

    lines.append("Projects")
    for proj in BASE_PROJECTS:
        header = f"{proj['title']} | {proj['context']} | {proj['dates']}"
        lines.append(header)
        for b in proj["bullets"]:
            lines.append(f"- {b}")
        lines.append("")

    lines.append("Education")
    for edu in BASE_EDUCATION:
        lines.append(edu["degree"])
        lines.append(f"{edu['institution']} | {edu['dates']} | {edu['grade']}")
        if edu.get("extra"):
            lines.append(edu["extra"])
        lines.append("")

    lines.append("Certifications")
    for c in BASE_CERTIFICATIONS:
        lines.append(f"- {c}")
    lines.append("")

    lines.append("Additional Information")
    for a in BASE_ADDITIONAL_INFO:
        lines.append(f"- {a}")

    return "\n".join(lines)


VISHNU_BASE_RESUME = build_base_resume_text()


# ═══════════════════════════════════════════════════════════════════
# 3. SAFETY: BANNED SKILLS (Vishnu does NOT have these at production level)
# Per intake form 9B. Used as a final scrub on tailored output.
# We don't fail — we just strip.
#
# Permissive carve-out: upskilling-roadmap items (PL-300, T-SQL Advanced,
# Tableau Public, AZ-900) are allowed ONLY if accompanied by the explicit
# string '(In Progress)' or '(in progress)'. See is_banned() for the logic.
# ═══════════════════════════════════════════════════════════════════

BANNED_SKILLS = {
    # Cloud at production level (intake form 9B: "AWS theoretical only")
    "aws", "amazon web services", "azure data factory", "azure synapse",
    "gcp", "google cloud platform", "google cloud",
    # Big data / ML platforms he doesn't have hands-on
    "databricks", "mlflow", "kubeflow", "sagemaker", "vertex ai",
    # MLOps / containerisation
    "docker", "kubernetes", "ci/cd pipelines for ml", "airflow",
    # Production LLM / generative AI
    "langchain", "llamaindex", "rag systems", "rag pipelines",
    "vector databases", "chroma", "pinecone", "faiss in production",
    "computer vision", "opencv production", "yolo", "object detection",
    # Production-only databases (he has coursework PostgreSQL/MySQL, not prod)
    "postgresql production", "mysql production", "mongodb",
    "redis", "cassandra", "dynamodb",
    # Streaming / big data
    "kafka", "spark", "pyspark", "hadoop", "hive", "flink",
    # Languages he doesn't claim
    "javascript", "typescript", "java production", "c++", "c#", "golang", "rust", "scala",
    # Frameworks he doesn't have
    "react", "angular", "vue", "node.js", "django production", "fastapi production",
    # Frontend / mobile
    "ios", "android", "swift", "kotlin", "flutter",
    # Certifications he hasn't earned (in-progress carve-out handled in is_banned)
    "pl-300 certified", "az-900 certified", "az-900 azure fundamentals certified",
    "dp-203 certified", "aws certified", "gcp certified",
    # SAP / Oracle / enterprise systems he has no hands-on
    "sap", "oracle erp", "salesforce", "servicenow administration",
    # QA / testing tools (unless he's on QA automation track)
    "selenium production", "cypress production", "playwright production",
    "appium", "testng",
    # Misc that frequently appear in JDs but he doesn't have
    "snowflake production", "snowpark", "dbt production",
    "tableau certified", "looker", "qlik", "qlikview",
}

# Strings that, if present in a skill mention, allow it through the banned filter.
# This is what enables the "permissive in-progress" mode for upskilling items.
IN_PROGRESS_MARKERS = (
    "(in progress)",
    "(in-progress)",
    "(currently learning)",
    "(currently studying)",
    "(target",
    "(targeted",
    "(planned",
)


# ═══════════════════════════════════════════════════════════════════
# 4. PROMPT — single, focused, no over-engineering
# ═══════════════════════════════════════════════════════════════════

ASTRA_PROMPT = """You are Astra, a resume tailoring engine for Vishnunath Nharekkat, a Data and Analytics professional based in Dublin, Ireland.

Your job: take Vishnu's BASE RESUME and the JOB DESCRIPTION, and produce a tailored version that mirrors the JD's language and priorities, while staying 100% truthful.

═══ HARD RULE: NO EM DASHES ANYWHERE IN YOUR OUTPUT ═══

This is non-negotiable. Em dashes (—, U+2014) are a leading recruiter signal of AI-generated writing in 2026. They MUST NOT appear in:
- The summary
- Any experience bullet
- Any project bullet
- The candidate_title
- The jd_fit_warning
- Any other field

Replace every em dash with one of: a comma, a full stop with sentence break, a colon, parentheses, or simply rephrase the sentence to not need one. Hyphens inside compound words (e.g. "cross-functional", "end-to-end") and en dashes inside date ranges (e.g. "June 2021 - September 2023") are fine and not affected by this rule.

Examples of em-dash phrasing to AVOID and how to fix:
- BAD:  "He is on Stamp 1G — CSEP-eligible."
  GOOD: "He is on Stamp 1G, CSEP-eligible."
- BAD:  "MSc Data Analytics — NFQ Level 9, NCI"
  GOOD: "MSc Data Analytics, NFQ Level 9, NCI"

═══ HARD RULE: NO PRONOUNS IN THE SUMMARY ═══

The professional profile / summary section MUST be written in pronoun-free, implied-subject style. This is the dominant Irish/UK CV convention.

DO NOT use any of these:
- First-person: "I", "I am", "I bring", "my experience"
- Third-person: "He", "He brings", "He is", "his experience", "Vishnu is", "Vishnu brings"
- Possessives: "my MSc", "his thesis"

INSTEAD: Drop the subject entirely. The reader knows it's about the candidate.

Examples of bad and good summary openings:
- BAD:  "I am a data analyst with an MSc..."
- BAD:  "He brings 2 years of engineering experience..."
- BAD:  "Vishnu holds an MSc in Data Analytics..."
- GOOD: "Data and Analytics professional with an MSc..."
- GOOD: "Brings 2 years of engineering experience reframed as quantitative work..."  (verb-led, no pronoun)
- GOOD: "Holds an MSc in Data Analytics from the National College of Ireland..."

This applies to the summary ONLY. Cover letters use first-person ("I", "my") which is normal letter convention. Bullets use action verbs (Built, Validated, Designed) which are already pronoun-free.

═══ CANDIDATE FACTS YOU CANNOT INVENT ═══

Vishnu's REAL skills (safe to claim, production-grade):
- Languages: Python (pandas, NumPy, scikit-learn, matplotlib, seaborn), SQL (intermediate, coursework + IMMCO Excel-based queries), R (basic, coursework).
- BI / Visualisation: Power BI (intermediate, including DAX and Power Query M, self-built dashboard), Microsoft Excel (advanced including PivotTables, VLOOKUP, XLOOKUP, INDEX/MATCH, used daily at IMMCO for KPI tracking).
- Machine Learning / Statistics: Regression, classification, clustering, feature engineering, hyperparameter tuning, model evaluation, hypothesis testing, A/B testing fundamentals, NLP, TF-IDF, LSTM (MSc thesis).
- ML frameworks at coursework/thesis level: scikit-learn, TensorFlow / Keras, NLTK (used in thesis), PyTorch (coursework familiarity only, do NOT claim production).
- Data Processes: Data cleansing, data validation, data quality assessment, ETL fundamentals, KPI reporting, dashboard development, star schema, data modelling.
- Tools: Git, GitHub, Jupyter, JIRA, Microsoft Office, Confluence.
- Geospatial / Engineering (legacy from IMMCO): QGIS, MapInfo, AutoCAD.
- Methods: Stakeholder management, cross-functional collaboration, requirements gathering, data storytelling, problem solving.
- Education: MSc Data Analytics (2:1, National College of Ireland, NFQ Level 9), Apr 2025. BSc Electronics (First Class, University of Calicut), 2020.
- Certifications: Power BI Data Analytics (Codebasics), SQL Basic (HackerRank), Excel (Codebasics), Python/DS/ML/AI/Power BI (Luminar Technolab), Python 101 for Data Science (IBM Cognitive Class).
- Currently Learning (in progress, NOT YET CERTIFIED): PL-300, T-SQL Advanced, Tableau Public, AZ-900.
- Work Authorisation: Stamp 1G (valid through April 2027). Eligible for Critical Skills Employment Permit at €40,904 salary floor (March 2026 threshold).

Vishnu does NOT have, and you must NEVER claim he does (intake form 9B):
- AWS / Azure / GCP at production level (theoretical/coursework only).
- Databricks, MLflow, Kubeflow, SageMaker, Vertex AI.
- Docker / Kubernetes for ML deployment.
- LangChain, RAG systems in production, vector databases at production level.
- Computer Vision, OpenCV, YOLO, object detection.
- Production PostgreSQL / MySQL / MongoDB (coursework familiarity only).
- Kafka, Spark, PySpark, Hadoop, Hive.
- JavaScript, TypeScript, Java, C++, C#, Go, Rust, Scala.
- React, Angular, Vue, Node.js.
- PL-300 / AZ-900 / DP-203 / AWS / GCP certifications (CURRENTLY LEARNING, NOT CERTIFIED).
- SAP, Oracle ERP, Salesforce administration.
- Snowflake / dbt / Looker / Qlik at production level.
- Hands-on Selenium / Cypress / Playwright (only relevant if QA automation archetype, and even then keep claims to coursework / planned).
- Tableau certification (uses Tableau Public at the in-progress level only).

If the JD demands any of the above as must-have, do NOT add them. Honesty is non-negotiable.

═══ STEP 1: DETECT THE JD ARCHETYPE ═══

Before tailoring, classify the JD into ONE of these eight archetypes. This drives how you frame the summary, bullets, and skills.

A. data_analyst
   Signals: titles like "Data Analyst", "Business Analyst", "Reporting Analyst", "Operations Analyst", "Junior Data Analyst", "Insights Analyst". Asks for SQL, Excel, Power BI or Tableau, data cleansing, KPI reporting, stakeholder reporting, 0-3 years experience. Often Microsoft-stack (Irish banks, ESB, DAA, AIB, BoI).
   Vishnu is a STRONG match. Lead with MSc Data Analytics, Python/SQL/Power BI stack, and the IMMCO experience reframed as quantitative validation and KPI tracking. The Power BI dashboard project is the headline portfolio piece.

B. bi_analyst
   Signals: "BI Analyst", "Business Intelligence Analyst", "Reporting Analyst with Power BI", "BI Developer". Heavy emphasis on Power BI or Tableau, DAX, dashboard development, semantic models, data modelling.
   Vishnu is a STRONG match. Lead with Power BI mastery (DAX, Power Query M, self-built dashboard), Excel depth, and IMMCO KPI tracking experience. Mention SQL data modelling.

C. junior_data_scientist
   Signals: "Junior Data Scientist", "Data Scientist", "AI Engineer", "Machine Learning Engineer", "ML Engineer". Asks for Python, scikit-learn, PyTorch / TensorFlow, NLP, deep learning, sometimes LLM/RAG, sometimes MLOps. 1-3 years often listed but graduate-friendly at junior level.
   Vishnu is a MODERATE match. Lead with the MSc thesis (TF-IDF + SVM vs LSTM, NLP, 87% accuracy) as the headline portfolio piece. Frame Python and scikit-learn as core. Be honest about production cloud / MLOps gaps. If the JD requires production LLM/RAG/vector-db work, set archetype to unknown_stretch instead.

D. analytics_engineer
   Signals: "Analytics Engineer", "Data Engineer (Junior)", "BI Engineer", titles asking for SQL + dbt + Snowflake + warehousing + data modelling. Often in fintech, SaaS, consulting.
   Vishnu is a MODERATE match. Lead with SQL data modelling, star schema (Power BI project), ETL fundamentals, and Python data pipelines. Be honest that dbt and Snowflake are coursework/planned, not production. Add jd_fit_warning if the JD demands 2+ years of dbt or Snowflake hands-on.

E. graduate_programme_tech
   Signals: "Graduate Programme" / "Graduate Scheme" / "Early Careers" PLUS a tech / data / analytics focus. Examples: Accenture Data & AI FY27, Mastercard Launch, Vodafone Discover Data Science & Analytics, Workday Generation, Salesforce Futureforce, Liberty IT TechStart.
   Vishnu is a STRONG match. Lead with MSc Data Analytics (NFQ Level 9), the thesis, the Power BI portfolio project. Frame IMMCO as "2 years and 3 months of engineering experience that gives me an operational lens most graduate applicants lack". Highlight Stamp 1G eligibility prominently because graduate programmes often filter on right-to-work.

F. graduate_programme_business
   Signals: "Graduate Programme" / "Business & Commercial Graduate" / "Management Graduate" / Big-4 graduate programmes (PwC, Deloitte, EY, KPMG) / consulting graduate schemes. Often "any quantitative degree" / "all STEM disciplines".
   Vishnu is a MODERATE-to-STRONG match. Lead with the MSc and the analytical mindset. Frame IMMCO as "operational and stakeholder experience inside a real engineering programme delivering FTTP networks across 250+ towns". Use business-strategic language, less heavy on Power BI jargon. The BSc Electronics First Class is evidence of quantitative excellence.

G. qa_automation
   Signals: "QA Automation Engineer", "Software Engineer in Test", "SDET", "Test Automation Engineer". Asks for Selenium / Playwright / Cypress, Python or Java test automation, CI/CD, JIRA, API testing.
   Vishnu is a WEAK match. He has QA validation experience at IMMCO (FTTP design QA, 20% rework reduction) but no software test automation tooling. Frame the IMMCO QA experience as "structured QA / validation discipline transferable to software test environments". Add a jd_fit_warning recommending Data QA / ETL Tester / BI Test Analyst hybrid roles where his profile is stronger. Still produce a usable resume.

H. unknown_stretch
   Signals: domain Vishnu has no experience in (construction, finance trading, niche tech) OR seniority well above his level OR a JD demanding production MLOps / LLM / cloud / Spark / Kafka as must-have.
   Soften the framing. Lead with transferable skills (data validation, KPI tracking, stakeholder management, MSc Data Analytics). Add a jd_fit_warning. Still produce the resume so he can decide whether to apply.

═══ STEP 2: TAILOR USING THE ARCHETYPE ═══

Once you've picked an archetype, follow the framing rules above when writing the summary, bullets, project bullets, and skills.

═══ OUTPUT, RETURN A JSON OBJECT WITH THIS EXACT SHAPE ═══

{
  "jd_archetype": "<one of: data_analyst | bi_analyst | junior_data_scientist | analytics_engineer | graduate_programme_tech | graduate_programme_business | qa_automation | unknown_stretch>",

  "jd_fit_warning": "<empty string if Vishnu is a strong match. Otherwise 1-2 short sentences explaining the mismatch and what alternative role at this company he could pursue. Examples: 'This JD requires production MLOps and LangChain RAG systems, which Vishnu has not built in production. Consider applying to the Data Analyst or BI Analyst roles at this company instead.' OR 'This is a senior data scientist role asking for 5+ years experience; Vishnu is at the graduate/0-2 year level and would be stronger applying for the Junior Data Scientist programme at this company.'>",

  "candidate_title": "<job-title-style line directly under the name. Mirror the JD's role title where possible. Examples: 'Data Analyst', 'BI Analyst', 'Junior Data Scientist', 'Analytics Engineer', 'Graduate Data Analyst', 'Business & Commercial Graduate', 'Data QA Analyst'. Default to 'Data & Analytics Professional' if unclear. Drop 'Senior'/'Lead'/'Principal' modifiers from the JD title since Vishnu is entry-to-junior level.>",

  "summary": "<EXACTLY 4-5 sentences. Natural flow, no choppy listing. NO em dashes. NO pronouns (no 'I', no 'he', no 'Vishnu', no 'his', no 'my'). Drop the subject entirely. The framing depends on the archetype:

  For archetypes A (data_analyst), B (bi_analyst):
    Sentence 1: Role identity matched to JD + 'MSc Data Analytics graduate (2:1, NCI)' + the 2 years 3 months engineering experience reframed as quantitative work. Pronoun-free.
       GOOD: 'Data Analyst with an MSc in Data Analytics (2:1, National College of Ireland) and 2 years and 3 months of engineering experience in quantitative validation, KPI reporting, and geospatial data analysis.'
    Sentence 2: Concrete technical stack relevant to the JD: Python (pandas, scikit-learn), SQL, Power BI / DAX / Power Query M, Excel.
       GOOD: 'Hands-on with Python (pandas, NumPy, scikit-learn), SQL, Power BI (DAX, Power Query M), and advanced Excel for data cleansing, dashboard development, and KPI reporting.'
    Sentence 3: One concrete portfolio metric: either the 20% rework reduction at IMMCO OR the Power BI dashboard project OR the 87% LSTM thesis (pick the one most relevant to the JD).
       GOOD: 'Reduced design errors and downstream rework by approximately 20 percent at IMMCO through structured QA and KPI tracking across N2P and RCP fibre programmes.'
    Sentence 4: Currently learning line if relevant to the JD (PL-300 for Power BI roles, AZ-900 for Azure-heavy roles, T-SQL Advanced for Microsoft-stack roles, Tableau Public for Tableau-heavy roles).
       GOOD: 'Currently working towards PL-300 (Power BI Data Analyst) and T-SQL Advanced certifications.'
    Sentence 5: Stamp 1G + CSEP-eligibility line.
       GOOD: 'Stamp 1G valid through April 2027 and eligible for a Critical Skills Employment Permit.'

  For archetype C (junior_data_scientist):
    Sentence 1: 'MSc Data Analytics graduate (2:1, NCI) with hands-on experience in NLP, classical and deep machine learning, and statistical analysis.' Pronoun-free.
    Sentence 2: MSc thesis as the headline: 'MSc thesis compared TF-IDF with Linear Support Vector Machine (86 percent accuracy) against a Long Short-Term Memory deep learning model (87 percent accuracy) on Indian political tweet sentiment, using Python, scikit-learn, NLTK, and TensorFlow / Keras.' Pronoun-free.
    Sentence 3: Other technical depth + tools relevant to the JD.
       GOOD: 'Comfortable with Python (pandas, NumPy, scikit-learn, matplotlib), SQL, feature engineering, hyperparameter tuning, and model evaluation including confusion matrices and F1 analysis.'
    Sentence 4: Optional, currently learning line if the JD mentions a specific tool.
    Sentence 5: Stamp 1G note.

  For archetype D (analytics_engineer):
    Sentence 1: 'Analytics Engineer with an MSc in Data Analytics (2:1, NCI) and a strong foundation in SQL data modelling, ETL design, and Power BI semantic modelling.' Pronoun-free.
    Sentence 2: Concrete stack: SQL (star schema design), Python (pandas, ETL scripts), Power BI (DAX, Power Query M).
    Sentence 3: Power BI portfolio project as the headline: 'Self-built end-to-end Power BI sales and operations dashboard with a SQL star-schema backend, DAX measures, and drill-through analytics.'
    Sentence 4: Currently learning line (T-SQL Advanced is the most relevant here).
    Sentence 5: Stamp 1G note.

  For archetypes E (graduate_programme_tech) and F (graduate_programme_business):
    Sentence 1: 'Recent MSc graduate in Data Analytics (2:1, NCI), NFQ Level 9, with prior 2 years and 3 months of engineering experience at IMMCO India.' Pronoun-free.
    Sentence 2: Verb-led description of the operational experience: 'Brings real operational depth and stakeholder-coordination experience from working on FTTP fibre rollout programmes alongside design, planning, and build teams.'
    Sentence 3: Headline portfolio asset relevant to the programme: the thesis (NLP, ML) for tech programmes, the Power BI dashboard for business / consulting programmes, the IMMCO 20% rework reduction for operations-flavoured programmes.
    Sentence 4: Why this programme. Connect to what the programme offers (rotation, breadth, leadership development). Mention the company by name if it appears clearly in the JD. Verb-led, no pronoun.
       GOOD: 'Drawn to the [Programme Name] for its rotation structure across data, analytics, and business strategy.'
    Sentence 5: Stamp 1G eligibility. Critical for graduate programmes that often state 'continuous right to work required'.

  For archetype G (qa_automation):
    Sentence 1: 'Data and QA professional with an MSc in Data Analytics (2:1, NCI) and 2 years and 3 months of QA and validation experience on enterprise FTTP fibre rollout programmes at IMMCO India.' Pronoun-free.
    Sentence 2: Concrete IMMCO QA metrics: 'Reduced design errors and downstream rework by approximately 20 percent through structured Quality Assurance and First Pass validation; owned QA workflow and onboarded new engineers into validation procedures.'
    Sentence 3: Technical stack: Python, SQL, Excel, Git, JIRA. Honest framing for test automation (planned / learning).
       GOOD: 'Comfortable with Python, SQL, Git, and JIRA; building familiarity with Selenium and API testing through self-directed learning.'
    Sentence 4: Stamp 1G note.

  For archetype H (unknown_stretch):
    Sentence 1: 'Data and Analytics professional with an MSc in Data Analytics (2:1, NCI)...' Pronoun-free.
    Sentence 2: Briefest IMMCO metric. Verb-led.
    Sentence 3: Education + portfolio.
    Sentence 4: 'Open to applying transferable analytical, data validation, and KPI reporting skills to new sectors.' + Stamp 1G note.

  Avoid robotic phrasing across all archetypes. No 'leveraging', 'utilizing', 'spearheading', 'passionate about', 'committed to excellence', 'seamless', 'innovative', 'cutting-edge', 'pivotal', 'tapestry'. Write like a confident human.>",

  "skills_additions": {
    "<existing category name>": ["<JD-relevant skill he actually has, OR an in-progress upskilling item with explicit '(In Progress)' marker>", "..."],
    "...": []
  },

  "experience_bullets": [
    {
      "company": "IMMCO",
      "responsibilities": ["<6 rewritten bullets>"],
      "achievements": ["<3 rewritten achievement bullets>"]
    }
  ],

  "project_bullets": [
    {
      "title": "Comparative Machine Learning for Political Tweet Sentiment Analysis",
      "bullets": ["<3 rewritten bullets>"]
    },
    {
      "title": "End-to-End Power BI Sales and Operations Dashboard",
      "bullets": ["<3 rewritten bullets>"]
    }
  ],

  "target_company": "<company name from JD, or 'Company' if not stated>"
}

═══ RULES FOR EACH FIELD ═══

SKILLS_ADDITIONS:
- The keys MUST be one of: "Languages and Querying", "Business Intelligence and Visualisation", "Machine Learning and Statistics", "Data Processes", "Tools and Productivity", "Geospatial and Engineering Tools", "Methods and Soft Skills".
- Add ONLY skills that appear in the JD AND that Vishnu genuinely has (per the "REAL skills" list above), OR explicit in-progress upskilling items with the marker '(In Progress)' attached.
- Examples of valid additions: "Microsoft Fabric (In Progress)", "Tableau Public (In Progress)", "T-SQL Advanced (In Progress)", "PL-300 (In Progress)", "AZ-900 (In Progress)", "Data Storytelling", "Requirements Gathering", "Reporting Automation".
- Do NOT add: AWS production, Azure production, GCP, Snowflake production, dbt production, Databricks, Spark, Kafka, MLflow, Docker, Kubernetes, LangChain, RAG production, Computer Vision, PostgreSQL production, MySQL production, MongoDB, Java, JavaScript, React, SAP, Oracle ERP. Even if the JD asks for them. Honesty is non-negotiable.
- If a category has no relevant additions, return an empty list for that category.
- 0-3 additions per category. Quality over quantity.

EXPERIENCE_BULLETS (IMMCO):
- Rewrite each bullet so its WORDING aligns with the JD's archetype, but every concrete claim must come from the BASE responsibilities/achievements provided.
- Every metric stays IDENTICAL: 20 percent rework reduction, 4 to 6 engineers, IDD/UDD/CP/ADD design phases, N2P/RCP programmes, FTTP/FTTH, SDU/MDU. Never change a number.
- IMMCO: keep 6 responsibility bullets and 3 achievement bullets.
- Each bullet starts with a strong past-tense verb: performed, ran, built, translated, acted, coordinated, reduced, standardised, validated, owned, mentored, tracked, reconciled.
- Do NOT introduce new tools, sectors, or claims not in the base bullets. You MAY re-label an existing claim using a JD-aligned synonym:
  - "validated data inconsistencies across AutoCAD, MapInfo, QGIS outputs" -> "performed data quality assessment across multiple geospatial data sources" (for data quality / data ops JDs)
  - "Excel KPI tracking workbooks" -> "KPI reporting and dashboard development in Excel" (for BI / reporting JDs)
  - "ran daily stand-ups" -> "led daily Agile stand-ups" (for Agile-flavoured JDs)
  - "FTTP and FTTH design" -> "telecommunications network design and rollout programmes" (for non-telecom audiences)
- NO em dashes inside bullets. Use commas, periods, colons, or parentheses.

PROJECT_BULLETS:
- Two projects, in this order: "Comparative Machine Learning for Political Tweet Sentiment Analysis" (thesis), "End-to-End Power BI Sales and Operations Dashboard" (self-directed).
- Each project keeps 3 bullets.
- Re-frame the wording to match the JD's archetype (thesis is the headline for junior_data_scientist; Power BI dashboard is the headline for data_analyst / bi_analyst / analytics_engineer).
- All metrics from the base bullets must be preserved EXACTLY: 86 percent SVM, 87 percent LSTM, 20,000 word dissertation, 2:1 grade.
- NO em dashes.

SUMMARY:
- 4-5 sentences. Count them.
- Mention the target company by name if it appears clearly in the JD.
- Include the Stamp 1G + CSEP-eligible note in the last sentence unless it would awkwardly displace something more important.
- One concrete metric from his real experience or thesis.
- Confident but not boastful. No hype words.
- NO em dashes. NO pronouns.

CANDIDATE_TITLE:
- Match the JD's role title verbatim where reasonable (max 8 words).
- Drop seniority modifiers ('Senior', 'Lead', 'Principal') if present in the JD title.

JD_ARCHETYPE & JD_FIT_WARNING:
- Always populate jd_archetype with one of the eight values listed.
- jd_fit_warning is empty string '' when archetype is A, B, E, F (strong fit). Populate it for C (if production cloud/MLOps required), D (if dbt/Snowflake production required), G, H with a concise honest assessment + alternative role suggestion.
- NO em dashes in the warning text either.

═══ FINAL CHECK BEFORE OUTPUT ═══
Scan your entire output for the em dash character (—). If you find any, remove or replace them. This applies to summary, bullets, project bullets, candidate_title, jd_fit_warning, and any other field.

═══ OUTPUT ═══
Return ONLY the JSON object. No prose, no markdown fences, no explanation.
"""


# Lighter cover letter prompt: single 4-paragraph structure, no war-story branching.
COVER_LETTER_PROMPT = """You are Vishnunath Nharekkat writing a cover letter for the role described below.
Write in first person. Sound like a real human, not a corporate template.

═══ HARD RULE: NO EM DASHES ANYWHERE IN THE LETTER ═══

Em dashes (—, U+2014) are a leading recruiter signal of AI-generated writing in 2026. Do NOT use them anywhere in the letter body. Use commas, periods, colons, or sentence breaks instead. Hyphens inside compound words ('cross-functional', 'end-to-end', 'first-hand') are fine. En dashes inside date ranges are fine. Em dashes are NOT.

Before you finalise, scan your letter for the em dash character (—) and remove every one.

═══ CONTEXT ═══
Vishnu is an entry-to-junior Data and Analytics professional based in Dublin.
- MSc Data Analytics, 2:1, National College of Ireland (Jan 2024 to Apr 2025). NFQ Level 9.
- 2 years and 3 months at IMMCO India as Telecom Engineer (Jun 2021 to Sep 2023): geospatial data analysis in QGIS / MapInfo, QA and First Pass validation on FTTP fibre design, KPI tracking in Excel, SME for 4 to 6 design engineers, 20 percent rework reduction across N2P and RCP fibre programmes.
- MSc thesis: TF-IDF + SVM (86 percent accuracy) vs LSTM (87 percent accuracy) on Indian political tweet sentiment, using Python, scikit-learn, NLTK, and TensorFlow / Keras.
- Self-built end-to-end Power BI dashboard with SQL star-schema backend, DAX measures, and drill-through analytics.
- BSc Electronics, First Class, University of Calicut (2017 to 2020).
- Stamp 1G valid through April 2027. Eligible for Critical Skills Employment Permit.
- Currently learning: PL-300, T-SQL Advanced, Tableau Public, AZ-900 (NOT YET CERTIFIED).

═══ HARD RULES ═══

DO NOT mention skills he does not have at production level:
- No AWS, Azure, GCP production. No Databricks, MLflow, Kubeflow, SageMaker, Vertex AI.
- No LangChain, RAG production, vector databases, Computer Vision.
- No Spark, Kafka, Hadoop, Hive, PySpark.
- No JavaScript, Java, C++, Go, Rust, Scala production.
- No PostgreSQL, MySQL, MongoDB production. No Snowflake production. No dbt production.
- No SAP, Oracle ERP, Salesforce administration.
- No PL-300, AZ-900, DP-203, AWS, GCP certifications (he is currently learning, NOT certified).
- No production Selenium / Cypress / Playwright.
- No Tableau certification (uses Tableau Public at the in-progress level).

DO NOT inflate experience:
- He has 2 years and 3 months of engineering experience plus the MSc. Not 3 years. Not 5.
- His commercial experience is engineering / QA / KPI reporting at IMMCO. If the JD is in banking, pharma, or another sector, FRAME his experience as transferable. Never claim sector-specific experience he does not have.

═══ BANNED PHRASES ═══
Do not use any of these (they make the letter sound AI-generated or template-y):
- "I am writing to express my interest"
- "I am excited to apply"
- "Please find my resume attached"
- "I believe I am a perfect fit"
- "passionate about", "driven by", "committed to excellence"
- "leveraging", "utilizing", "harnessing"
- "showcasing", "highlighting", "demonstrating"
- "testament to", "underscores", "pivotal", "tapestry", "realm"
- "seamless", "innovative", "groundbreaking", "cutting-edge"
- "at the forefront of", "at the intersection of"
- Three-adjective chains ("scalable, reliable, and efficient")
- ANY em dash (—)

═══ STRUCTURE ═══
4 short paragraphs, plain text, no markdown, no headers, no bold.

Paragraph 1 (Hook, 2-3 sentences):
Open by referring to a SPECIFIC analytical / technical / business challenge or focus from the JD, not the company in general. Show you actually read what they wrote. Mention the role title from the JD and the company name in this paragraph.

Examples of good hooks:
- "Building a Power BI semantic model that thousands of business users actually trust takes more than DAX skill. It depends on the data quality and the validation logic behind the scenes."
- "Translating raw data into KPI reporting that 250+ store managers act on is exactly the kind of problem I worked on every day in my final year at IMMCO India."
- "The [Programme Name] structure across data, analytics, and business strategy rotations is exactly the breadth I want in my first Irish data role."

Paragraph 2 (Connection, 3-4 sentences):
Connect Vishnu's MSc + IMMCO experience to the role. Pick the most JD-relevant of these three angles:
- Technical: Python / SQL / Power BI / DAX / scikit-learn matched to the JD's stack.
- Operational: KPI tracking, QA validation, 20 percent rework reduction, SME role for 4-6 engineers.
- Academic: MSc thesis (NLP, LSTM, 87 percent accuracy) or the Power BI portfolio project.
Use exact metrics. Never round or change numbers.

Paragraph 3 (Why this company / programme, 2 sentences):
Brief reason this specific company appeals. Connect to something from the JD (a product, a value, a programme structure). Avoid generic flattery.

Paragraph 4 (Close, 2-3 sentences):
Mention Stamp 1G eligibility plainly: "I hold a Stamp 1G permission valid through April 2027 and am eligible for a Critical Skills Employment Permit." Express interest in discussing the role. End with "Thank you," on its own line, then "Vishnunath Nharekkat" on the next line.

═══ STYLE ═══
- Vary sentence length. Mix short with longer.
- NO em dashes. Use commas, periods, colons, or rephrase.
- Use plain verbs: built, validated, designed, ran, coordinated, reduced, standardised.
- He is entry-to-junior level. Confident, not arrogant. Eager to learn, not desperate.
- Length: 240-340 words total in the letter body (excluding "Thank you, Vishnunath Nharekkat" sign-off).

═══ FINAL CHECK ═══
Before you output, scan your letter once more for the em dash character (—) and remove every instance.

═══ OUTPUT ═══
Return ONLY the letter body as plain text. No "Dear Hiring Manager" greeting (the renderer adds it). No subject line. No address blocks. No markdown. No bold. No code fences.
Start directly with paragraph 1.
End with "Thank you," on its own line and "Vishnunath Nharekkat" on the next line.
"""


# ═══════════════════════════════════════════════════════════════════
# 5. SCHEMA
# ═══════════════════════════════════════════════════════════════════

class ExperienceBullets(BaseModel):
    company: str
    responsibilities: List[str]
    achievements: List[str] = Field(default_factory=list)


class ProjectBullets(BaseModel):
    title: str
    bullets: List[str]


class TailoredOutput(BaseModel):
    candidate_title: str
    summary: str
    skills_additions: dict
    experience_bullets: List[ExperienceBullets]
    project_bullets: List[ProjectBullets]
    target_company: str = "Company"
    jd_archetype: str = ""
    jd_fit_warning: str = ""


# ═══════════════════════════════════════════════════════════════════
# 6. CALL GEMINI
# ═══════════════════════════════════════════════════════════════════

def call_gemini(api_key: str, jd_text: str, resume_text: str = "") -> dict:
    """Single API call. JSON mode. Returns parsed dict or {'error': ...}.
    resume_text is Vishnu's base resume, passed through so the model sees
    what he sees in the UI box and respects any edits he made there.
    Falls back to the constants-derived default if not supplied."""
    if not api_key:
        return {"error": "Missing GOOGLE_API_KEY."}
    if not jd_text or not jd_text.strip():
        return {"error": "Job description is empty."}

    if not resume_text or not resume_text.strip():
        resume_text = VISHNU_BASE_RESUME

    client = genai.Client(api_key=api_key)
    prompt = (
        f"{ASTRA_PROMPT}\n\n"
        f"═══ BASE RESUME ═══\n{resume_text}\n\n"
        f"═══ JOB DESCRIPTION ═══\n{jd_text}"
    )

    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.4,  # low for reliability
            ),
        )
        text = response.text.strip()
        # Strip code fences if model wrapped them despite JSON mode
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        return json.loads(text)
    except json.JSONDecodeError as e:
        return {"error": f"Could not parse JSON from model: {e}"}
    except Exception as e:
        return {"error": f"Generation error: {e}"}


def generate_cover_letter(api_key: str, resume_data: dict, jd_text: str) -> str:
    """Single API call. Plain text mode. Returns letter body or error message starting with 'ERROR:'."""
    if not api_key:
        return "ERROR: Missing GOOGLE_API_KEY."
    if not jd_text or not jd_text.strip():
        return "ERROR: Job description is empty."

    client = genai.Client(api_key=api_key)

    # Pass the tailored resume context so the cover letter aligns with what the JD-tailored CV says.
    resume_context = (
        f"Tailored role title: {resume_data.get('candidate_title', '')}\n"
        f"Target company: {resume_data.get('target_company', '')}\n"
        f"Tailored summary: {resume_data.get('summary', '')}"
    )

    prompt = (
        f"{COVER_LETTER_PROMPT}\n\n"
        f"═══ TAILORED RESUME CONTEXT ═══\n{resume_context}\n\n"
        f"═══ JOB DESCRIPTION ═══\n{jd_text}"
    )

    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.6),
        )
        text = (response.text or "").strip()
        # Strip code fences if any
        if text.startswith("```"):
            text = re.sub(r"^```(?:\w+)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        # Strip any "Dear ..." greeting if the model added one despite instructions
        text = re.sub(r"^dear\s+[^\n]+\n+", "", text, flags=re.IGNORECASE)
        # Final scrub
        text = scrub_banned_from_text(text)
        return text
    except Exception as e:
        return f"ERROR: Cover letter generation failed: {e}"


# ═══════════════════════════════════════════════════════════════════
# 7. ASSEMBLE: merge model output with base resume into final structure
# ═══════════════════════════════════════════════════════════════════

def has_in_progress_marker(text: str) -> bool:
    """True if the skill string contains an explicit in-progress marker.
    This is what gates the permissive carve-out for upskilling items."""
    t = text.lower()
    return any(marker in t for marker in IN_PROGRESS_MARKERS)


def is_banned(skill: str) -> bool:
    """Banned-skill check with permissive in-progress carve-out.
    A skill that matches a banned token is allowed through IF it carries
    an explicit '(In Progress)' or equivalent marker. This lets Astra add
    'PL-300 (In Progress)' or 'AZ-900 (In Progress)' even though the bare
    'PL-300 certified' / 'AZ-900 certified' tokens are banned."""
    s = skill.lower().strip()
    # If the skill has an in-progress marker, it's allowed even if a banned
    # token appears inside it. This handles "PL-300 (In Progress)" cleanly.
    if has_in_progress_marker(s):
        return False
    return any(b in s for b in BANNED_SKILLS)


def strip_em_dashes(text: str) -> str:
    """Replace em dashes (U+2014) with cleaner alternatives. Em dashes are a
    leading recruiter signal of AI-generated writing in 2026; we want zero in
    rendered output. This is the runtime safety net in case the model slips
    one through despite the prompt instructions.

    Replacement strategy:
    - " — " (em dash with spaces) -> ", " (comma + space) which reads naturally in 95% of cases
    - "word—word" (em dash without spaces) -> "word, word"

    We deliberately do NOT touch hyphens (-, U+002D) or en dashes (–, U+2013).
    Hyphens are correct in compound words (cross-functional). En dashes are
    correct in date ranges (June 2021 – September 2023). Neither is an AI signal.
    """
    if not text:
        return text
    cleaned = re.sub(r"\s*—\s*", ", ", text)
    cleaned = re.sub(r",\s*,", ",", cleaned)
    return cleaned


def strip_summary_pronouns(text: str) -> str:
    """Remove first-person/third-person pronouns from the start of summary
    sentences. Resume summaries follow the Irish/UK convention of pronoun-free,
    implied-subject prose ('Data Analyst with...', not 'I am a Data Analyst...'
    or 'He is a Data Analyst...').

    This is the runtime safety net. The prompt tells the model to write
    pronoun-free, but if a stray 'He brings...' or 'I bring...' or 'Vishnu is...'
    slips through, we catch it before rendering.

    Applied ONLY to the summary string. Cover letters use first-person which
    is normal letter convention and must NOT be stripped.
    """
    if not text:
        return text

    patterns = [
        # "He is a/an X" / "Vishnu is a/an X" / "I am a/an X" -> "X"
        (r"(^|(?<=\.\s))(I am|I'm|He is|He's|She is|She's|Vishnu is|Vishnu was|Vishnunath is|Vishnunath was)\s+(a|an)\s+", r"\1"),
        # "He is X" / "Vishnu is X" / "I am X" -> "X"
        (r"(^|(?<=\.\s))(I am|I'm|He is|He's|She is|She's|Vishnu is|Vishnu was|Vishnunath is|Vishnunath was)\s+", r"\1"),
        # "He brings/has/holds/builds X" -> "Brings/Has/Holds X" (verb stays, capitalised)
        (r"(^|(?<=\.\s))(He|She)\s+(brings|has|holds|builds|built|designs|designed|managed|manages|coordinates|coordinated|oversees|oversaw|maintains|maintained|trains|trained|reduced|reduces|improved|improves|streamlined|streamlines|validated|validates|tracked|tracks|delivered|delivers|leads|led|supports|supported|drives|drove|developed|develops|analyses|analysed|ran|runs)\b",
         lambda m: m.group(1) + m.group(3).capitalize()),
        # "I bring/have/hold X" -> "Bring/Have/Hold X"
        (r"(^|(?<=\.\s))I\s+(bring|have|hold|build|built|design|designed|manage|managed|coordinate|coordinated|oversee|oversaw|maintain|maintained|train|trained|reduce|reduced|improve|improved|streamline|streamlined|validate|validated|track|tracked|deliver|delivered|lead|led|support|supported|drive|drove|develop|developed|analyse|analysed|run|ran)\b",
         lambda m: m.group(1) + m.group(2).capitalize()),
        # "Vishnu/Vishnunath verb..." -> "Verb..."
        (r"(^|(?<=\.\s))(?:Vishnu|Vishnunath)\s+(built|designed|managed|coordinated|oversaw|maintained|trained|reduced|improved|streamlined|validated|tracked|delivered|led|supported|drove|developed|analysed|ran|brings|has|holds|builds|designs|manages|coordinates|oversees|maintains|trains|reduces|improves|streamlines|validates|tracks|delivers|leads|supports|drives|develops|analyses|runs)\b",
         lambda m: m.group(1) + m.group(2).capitalize()),
        # Possessives: "his MSc" / "my MSc" -> "MSc"
        (r"\b(?:his|His|My|my)\s+(MSc|BSc|thesis|experience|background|certification|portfolio)\b", r"\1"),
    ]

    cleaned = text
    for pattern, replacement in patterns:
        cleaned = re.sub(pattern, replacement, cleaned)

    # First letter of result should be uppercase
    if cleaned and cleaned[0].islower():
        cleaned = cleaned[0].upper() + cleaned[1:]

    # Tidy double spaces
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned


def scrub_banned_from_text(text: str) -> str:
    """Strip banned skill mentions and em dashes from a free-text string
    (summary, bullets, cover letter). Three-stage scrub:
    1. Remove banned skill tokens (AWS production, MLflow, Databricks, etc.)
       UNLESS the token is wrapped in an in-progress phrase. We can't easily
       parse in-progress markers from free text, so we apply a softer rule
       here: only strip the bare token if it's NOT followed by '(In Progress)'
       or similar.
    2. Em-dash strip.
    3. Tidy whitespace.
    """
    cleaned = text
    for b in BANNED_SKILLS:
        # Match the banned token followed optionally by an in-progress marker.
        # If the marker is present, leave it alone. Otherwise, strip.
        pattern = re.compile(
            r"\b" + re.escape(b) + r"\b(?!\s*\((?:in progress|in-progress|currently learning|currently studying|target|targeted|planned)\b)[, ]*",
            re.IGNORECASE,
        )
        cleaned = pattern.sub("", cleaned)
    # Tidy
    cleaned = re.sub(r",\s*,", ",", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    cleaned = re.sub(r"\s+([,.])", r"\1", cleaned)
    cleaned = strip_em_dashes(cleaned)
    return cleaned.strip()


def _normalize_skill(s: str) -> str:
    """Normalise a skill for dedup comparison. Strips parens, punctuation,
    and lowercases. So 'Power BI (DAX, Power Query M)' and 'Power BI DAX'
    both contain the normalised core 'power bi dax'."""
    if not s:
        return ""
    n = s.lower()
    # Strip in-progress markers before normalising so the comparison ignores them
    for marker in IN_PROGRESS_MARKERS:
        n = n.replace(marker, " ")
    n = re.sub(r"[()]", " ", n)
    n = re.sub(r"[^a-z0-9& ]", " ", n)
    n = re.sub(r"\s+", " ", n).strip()
    return n


def _is_duplicate_skill(candidate: str, existing_normalized: set) -> bool:
    """Return True if `candidate` is a duplicate of a skill already in
    `existing_normalized`. Handles exact match, token-subset, and token-superset
    cases. The substring check uses full-token containment to avoid false
    positives ('Modelling' vs 'Data Modelling' are NOT duplicates).
    """
    cand_norm = _normalize_skill(candidate)
    if not cand_norm:
        return True
    if cand_norm in existing_normalized:
        return True

    cand_tokens = set(cand_norm.split())
    for existing_norm in existing_normalized:
        existing_tokens = set(existing_norm.split())
        if not existing_tokens or not cand_tokens:
            continue
        if cand_tokens.issubset(existing_tokens):
            return True
        if existing_tokens.issubset(cand_tokens):
            return True
    return False


def merge_skills(additions: dict) -> dict:
    """Combine BASE_SKILLS with model additions. Drops banned terms and
    cross-category duplicates so the rendered skills section is clean."""
    final = {}
    seen_normalized = set()
    for category, base_list in BASE_SKILLS.items():
        for skill in base_list:
            seen_normalized.add(_normalize_skill(skill))

    for category, base_list in BASE_SKILLS.items():
        final[category] = list(base_list)

    if not isinstance(additions, dict):
        additions = {}

    for category, added_list in additions.items():
        if category not in final:
            continue
        if not isinstance(added_list, list):
            continue
        for skill in added_list:
            if not isinstance(skill, str):
                continue
            s = skill.strip()
            if not s:
                continue
            if is_banned(s):
                continue
            if _is_duplicate_skill(s, seen_normalized):
                continue
            final[category].append(s)
            seen_normalized.add(_normalize_skill(s))
    return final


def merge_experience(model_bullets: list) -> list:
    """For each base role, replace bullets with the model's tailored versions
    (after scrubbing). If the model dropped a role or returned the wrong count,
    fall back to base bullets so the resume stays complete."""
    by_company = {}
    if isinstance(model_bullets, list):
        for item in model_bullets:
            if isinstance(item, dict) and item.get("company"):
                by_company[item["company"]] = item

    final = []
    for base_role in BASE_EXPERIENCE:
        company = base_role["company"]
        tailored = by_company.get(company, {})
        tailored_resps = tailored.get("responsibilities") or []
        tailored_achs = tailored.get("achievements") or []

        clean_resps = [scrub_banned_from_text(r) for r in tailored_resps if isinstance(r, str)]
        clean_resps = [r for r in clean_resps if r.strip()]

        clean_achs = [scrub_banned_from_text(a) for a in tailored_achs if isinstance(a, str)]
        clean_achs = [a for a in clean_achs if a.strip()]

        expected_resp_count = len(base_role["responsibilities"])
        expected_ach_count = len(base_role["achievements"])

        if len(clean_resps) < max(1, expected_resp_count - 1):
            clean_resps = list(base_role["responsibilities"])
        if expected_ach_count > 0 and len(clean_achs) < expected_ach_count:
            clean_achs = list(base_role["achievements"])

        final.append({
            "role_title": base_role["role_title"],
            "company": base_role["company"],
            "location": base_role["location"],
            "dates": base_role["dates"],
            "responsibilities": clean_resps,
            "achievements": clean_achs,
        })
    return final


def merge_projects(model_projects: list) -> list:
    """For each base project, replace bullets with the model's tailored versions
    (after scrubbing). Falls back to base if model returned malformed data."""
    by_title = {}
    if isinstance(model_projects, list):
        for item in model_projects:
            if isinstance(item, dict) and item.get("title"):
                # Title match is loose: lowercase + prefix-match to handle minor variations
                by_title[item["title"].strip().lower()] = item

    final = []
    for base_proj in BASE_PROJECTS:
        base_title_key = base_proj["title"].strip().lower()
        tailored = None
        # Try exact match first, then loose prefix match
        if base_title_key in by_title:
            tailored = by_title[base_title_key]
        else:
            for k, v in by_title.items():
                if k.startswith(base_title_key[:30]) or base_title_key.startswith(k[:30]):
                    tailored = v
                    break

        tailored_bullets = (tailored or {}).get("bullets") or []
        clean_bullets = [scrub_banned_from_text(b) for b in tailored_bullets if isinstance(b, str)]
        clean_bullets = [b for b in clean_bullets if b.strip()]

        expected_count = len(base_proj["bullets"])
        if len(clean_bullets) < max(1, expected_count - 1):
            clean_bullets = list(base_proj["bullets"])

        final.append({
            "title": base_proj["title"],
            "context": base_proj["context"],
            "dates": base_proj["dates"],
            "bullets": clean_bullets,
        })
    return final


def assemble_resume(model_output: dict) -> dict:
    """Combine model output + base facts into a fully-populated resume dict."""
    summary = scrub_banned_from_text(model_output.get("summary", "") or "")
    summary = strip_summary_pronouns(summary)
    if not summary or len(summary.split()) < 20:
        # Fallback summary if model failed
        summary = (
            "MSc Data Analytics graduate (2:1, National College of Ireland) with 2 years "
            "and 3 months of engineering experience in quantitative validation, geospatial "
            "data analysis, and KPI-led programme tracking. Hands-on with Python (pandas, "
            "NumPy, scikit-learn), SQL, Power BI, and advanced Excel for data cleansing, "
            "dashboard development, and statistical analysis. MSc thesis applied NLP and "
            "deep learning (TF-IDF with SVM, and LSTM) to political tweet sentiment "
            "classification, with the LSTM model reaching 87 percent accuracy. Open to "
            "Data Analyst, Business Intelligence, and Analytics Engineer roles across "
            "Ireland. Stamp 1G valid through April 2027 and eligible for a Critical "
            "Skills Employment Permit."
        )

    return {
        "candidate_name": CANDIDATE_NAME,
        "candidate_title": strip_em_dashes((model_output.get("candidate_title") or CANDIDATE_TAGLINE).strip()),
        "contact_info": CANDIDATE_CONTACT,
        "summary": summary,
        "skills": merge_skills(model_output.get("skills_additions", {}) or {}),
        "experience": merge_experience(model_output.get("experience_bullets", []) or []),
        "projects": merge_projects(model_output.get("project_bullets", []) or []),
        "education": list(BASE_EDUCATION),
        "certifications": list(BASE_CERTIFICATIONS),
        "additional_info": list(BASE_ADDITIONAL_INFO),
        "target_company": strip_em_dashes((model_output.get("target_company") or "Company").strip()),
        "jd_archetype": (model_output.get("jd_archetype") or "").strip(),
        "jd_fit_warning": strip_em_dashes((model_output.get("jd_fit_warning") or "").strip()),
    }


# ═══════════════════════════════════════════════════════════════════
# 8. DOCX RENDERER — clean, ATS-friendly, single-column, no tables
# Matches the Astra Resume Formatting Playbook (US Letter, Times New Roman,
# 0.5" margins, 28pt name, 14pt title, 12pt body throughout).
# ═══════════════════════════════════════════════════════════════════

def _set_font(run, size, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        pass


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(text), 12, bold=True)


def _add_bullet(doc, text, bold_prefix=None, justify=True):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY if justify else WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(0)
    if bold_prefix:
        _set_font(p.add_run(bold_prefix), 12, bold=True)
    _set_font(p.add_run(text), 12)


def render_docx(data: dict) -> bytes:
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)

    # ─── Header ───
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(data["candidate_name"])
    run.font.all_caps = True
    _set_font(run, 28, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run(data["candidate_title"]), 14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(data["contact_info"]), 12, bold=True)

    # ─── Professional Profile ───
    _add_section_heading(doc, "Professional Profile")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run(data["summary"]), 12)

    # ─── Skills ───
    _add_section_heading(doc, "Key Skills, Tools and Technologies")
    for cat, skill_list in data["skills"].items():
        _add_bullet(doc, ", ".join(skill_list), bold_prefix=f"{cat}: ")

    # ─── Experience ───
    _add_section_heading(doc, "Professional Experience")
    for role in data["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        _set_font(p.add_run(header), 12, bold=True)

        for resp in role["responsibilities"]:
            _add_bullet(doc, resp)

        if role["achievements"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            _set_font(p.add_run("Achievements:"), 12, bold=True)
            for ach in role["achievements"]:
                _add_bullet(doc, ach)

    # ─── Projects ───
    _add_section_heading(doc, "Projects")
    for proj in data["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        header = f"{proj['title']} | {proj['context']} | {proj['dates']}"
        _set_font(p.add_run(header), 12, bold=True)
        for b in proj["bullets"]:
            _add_bullet(doc, b)

    # ─── Education ───
    _add_section_heading(doc, "Education")
    for edu in data["education"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        _set_font(p.add_run(edu["degree"]), 12, bold=True)

        meta = f"{edu['institution']} | {edu['dates']} | {edu['grade']}"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _set_font(p.add_run(meta), 12)

        if edu.get("extra"):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.paragraph_format.space_after = Pt(0)
            _set_font(p.add_run(edu["extra"]), 12)

    # ─── Certifications ───
    _add_section_heading(doc, "Certifications")
    for cert in data["certifications"]:
        _add_bullet(doc, cert)

    # ─── Additional Information ───
    _add_section_heading(doc, "Additional Information")
    for line in data["additional_info"]:
        _add_bullet(doc, line)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 9. PDF RENDERER — same content, ATS-friendly, single column
# ═══════════════════════════════════════════════════════════════════

def _esc(t):
    if t is None:
        return ""
    return escape(str(t))


def render_pdf(data: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()

    sn = ParagraphStyle("N", parent=styles["Normal"], fontName="Times-Roman",
                        fontSize=12, leading=14, alignment=TA_JUSTIFY, spaceAfter=0)
    sn_left = ParagraphStyle("NL", parent=styles["Normal"], fontName="Times-Roman",
                             fontSize=12, leading=14, alignment=TA_LEFT, spaceAfter=0)
    sh_name = ParagraphStyle("HN", parent=styles["Normal"], fontName="Times-Bold",
                             fontSize=22, leading=26, alignment=TA_CENTER, spaceAfter=0)
    sh_title = ParagraphStyle("HT", parent=styles["Normal"], fontName="Times-Bold",
                              fontSize=14, leading=16, alignment=TA_CENTER, spaceAfter=0)
    sh_contact = ParagraphStyle("HC", parent=styles["Normal"], fontName="Times-Bold",
                                fontSize=12, leading=14, alignment=TA_CENTER, spaceAfter=6)
    s_sec = ParagraphStyle("Sec", parent=styles["Normal"], fontName="Times-Bold",
                           fontSize=12, leading=14, alignment=TA_LEFT,
                           spaceBefore=10, spaceAfter=2)

    el = []

    # Header
    el.append(Paragraph(_esc(data["candidate_name"]).upper(), sh_name))
    el.append(Paragraph(_esc(data["candidate_title"]), sh_title))
    el.append(Paragraph(_esc(data["contact_info"]), sh_contact))

    # Summary
    el.append(Paragraph("Professional Profile", s_sec))
    el.append(Paragraph(_esc(data["summary"]), sn))

    # Skills
    el.append(Paragraph("Key Skills, Tools and Technologies", s_sec))
    skill_items = []
    for cat, skill_list in data["skills"].items():
        text = f"<b>{_esc(cat)}:</b> {_esc(', '.join(skill_list))}"
        skill_items.append(ListItem(Paragraph(text, sn_left), leftIndent=0))
    if skill_items:
        el.append(ListFlowable(skill_items, bulletType="bullet",
                               start="\u2022", leftIndent=15))

    # Experience
    el.append(Paragraph("Professional Experience", s_sec))
    for role in data["experience"]:
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        el.append(Paragraph(f"<b>{_esc(header)}</b>", sn_left))
        el.append(Spacer(1, 2))
        items = [ListItem(Paragraph(_esc(r), sn), leftIndent=0)
                 for r in role["responsibilities"] if r.strip()]
        if items:
            el.append(ListFlowable(items, bulletType="bullet", start="\u2022", leftIndent=15))
        if role["achievements"]:
            el.append(Paragraph("<b>Achievements:</b>", sn_left))
            ach_items = [ListItem(Paragraph(_esc(a), sn), leftIndent=0)
                         for a in role["achievements"] if a.strip()]
            if ach_items:
                el.append(ListFlowable(ach_items, bulletType="bullet",
                                       start="\u2022", leftIndent=25))
        el.append(Spacer(1, 4))

    # Projects
    el.append(Paragraph("Projects", s_sec))
    for proj in data["projects"]:
        header = f"{proj['title']} | {proj['context']} | {proj['dates']}"
        el.append(Paragraph(f"<b>{_esc(header)}</b>", sn_left))
        el.append(Spacer(1, 2))
        items = [ListItem(Paragraph(_esc(b), sn), leftIndent=0)
                 for b in proj["bullets"] if b.strip()]
        if items:
            el.append(ListFlowable(items, bulletType="bullet", start="\u2022", leftIndent=15))
        el.append(Spacer(1, 4))

    # Education
    el.append(Paragraph("Education", s_sec))
    for edu in data["education"]:
        el.append(Paragraph(f"<b>{_esc(edu['degree'])}</b>", sn_left))
        meta = f"{edu['institution']} | {edu['dates']} | {edu['grade']}"
        el.append(Paragraph(_esc(meta), sn_left))
        if edu.get("extra"):
            el.append(Paragraph(_esc(edu["extra"]), sn))
        el.append(Spacer(1, 4))

    # Certifications
    el.append(Paragraph("Certifications", s_sec))
    cert_items = [ListItem(Paragraph(_esc(c), sn_left), leftIndent=0)
                  for c in data["certifications"]]
    if cert_items:
        el.append(ListFlowable(cert_items, bulletType="bullet",
                               start="\u2022", leftIndent=15))

    # Additional
    el.append(Paragraph("Additional Information", s_sec))
    add_items = [ListItem(Paragraph(_esc(a), sn_left), leftIndent=0)
                 for a in data["additional_info"]]
    if add_items:
        el.append(ListFlowable(add_items, bulletType="bullet",
                               start="\u2022", leftIndent=15))

    doc.build(el)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 9b. COVER LETTER DOCX RENDERER
# ═══════════════════════════════════════════════════════════════════

def render_cover_letter_docx(letter_body: str, target_company: str = "") -> bytes:
    """Render a cover letter as a clean DOCX. Body is passed in as plain text;
    the renderer adds the candidate header, date, greeting, and sign-off."""
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1.0)

    def _line(text, bold=False, space_after=0, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            _set_font(p.add_run(text), 11, bold=bold)
        return p

    # ─── Sender header ───
    _line(CANDIDATE_NAME, bold=True, space_after=0)
    for piece in [c.strip() for c in CANDIDATE_CONTACT.split("|")]:
        if piece:
            _line(piece, space_after=0)

    _line("", space_after=10)

    # ─── Date ───
    today = datetime.date.today().strftime("%d %B %Y")
    _line(today, space_after=14)

    # ─── Greeting ───
    if target_company.strip():
        greeting = f"Dear Hiring Team at {target_company.strip()},"
    else:
        greeting = "Dear Hiring Team,"
    _line(greeting, space_after=10)

    # ─── Body paragraphs ───
    body = letter_body.strip()
    sign_off_pattern = re.compile(
        r"\n+(thank you,?|regards,?|sincerely,?|kind regards,?|best regards,?)\s*\n+vishnunath nharekkat\s*$",
        re.IGNORECASE,
    )
    sign_off_match = sign_off_pattern.search(body)
    if sign_off_match:
        body_main = body[: sign_off_match.start()].strip()
        sign_word = sign_off_match.group(1).rstrip(",").strip().capitalize()
    else:
        body_main = body
        sign_word = "Thank you"

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body_main) if p.strip()]
    if not paragraphs:
        paragraphs = [p.strip() for p in body_main.split("\n") if p.strip()]

    for para in paragraphs:
        clean_para = re.sub(r"\s*\n\s*", " ", para)
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        _set_font(p.add_run(clean_para), 11)

    # ─── Sign-off ───
    _line("", space_after=6)
    _line(f"{sign_word},", space_after=18)
    _line(CANDIDATE_NAME, space_after=0)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 10. STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════

st.set_page_config(page_title=PAGE_TITLE, layout="wide", page_icon="📊",
                   initial_sidebar_state="expanded")

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .block-container {padding-top: 1.5rem;}
    div.stButton > button:first-child {border-radius: 6px; font-weight: 600;}
</style>
""", unsafe_allow_html=True)

# Session state
if "tailored" not in st.session_state:
    st.session_state["tailored"] = None
if "saved_jd" not in st.session_state:
    st.session_state["saved_jd"] = ""
if "saved_base" not in st.session_state:
    st.session_state["saved_base"] = VISHNU_BASE_RESUME
if "cover_letter" not in st.session_state:
    st.session_state["cover_letter"] = None

# Sidebar
with st.sidebar:
    st.header("⚙️ Configuration")
    if GOOGLE_API_KEY:
        st.success("API key configured")
        api_key = GOOGLE_API_KEY
    else:
        st.warning("Add GOOGLE_API_KEY to Streamlit secrets, or paste it below.")
        api_key = st.text_input("Google API Key", type="password")

    st.divider()
    st.markdown("**Target roles:**")
    st.caption(
        "**Direct hire:** Data Analyst · BI Analyst · Junior Data Scientist · "
        "Analytics Engineer · Data QA Analyst\n\n"
        "**Graduate programmes:** Accenture Data & AI · Mastercard Launch · "
        "Vodafone Discover · Workday Generation · Salesforce Futureforce · "
        "Liberty IT TechStart · Big 4"
    )

    st.divider()
    st.markdown(f"**Model:** `{GENERATION_MODEL}`")

    st.divider()
    if st.button("🗑️ Reset", use_container_width=True):
        st.session_state["tailored"] = None
        st.session_state["saved_jd"] = ""
        st.session_state["saved_base"] = VISHNU_BASE_RESUME
        st.session_state["cover_letter"] = None
        st.rerun()

    st.caption("Astra v1.0 — Vishnu")

# Main UI
if not st.session_state["tailored"]:
    st.markdown(f"<h1 style='text-align: center;'>{PAGE_TITLE}</h1>", unsafe_allow_html=True)
    st.markdown(
        "<p style='text-align: center; color: #888;'>"
        "Paste a job description. Get a tailored, ATS-friendly resume.</p>",
        unsafe_allow_html=True,
    )
    st.divider()

    col_resume, col_jd = st.columns(2)

    with col_resume:
        st.subheader("📋 Base Resume")
        st.caption("Pre-filled with Vishnu's profile. Edit only if you need a one-off tweak for this application.")
        base = st.text_area(
            "Base Resume",
            value=st.session_state["saved_base"],
            height=420,
            label_visibility="collapsed",
        )
        if st.button("↩️ Restore default base resume", use_container_width=True):
            st.session_state["saved_base"] = VISHNU_BASE_RESUME
            st.rerun()

    with col_jd:
        st.subheader("💼 Job Description")
        st.caption("Paste the full JD here.")
        jd = st.text_area(
            "Job Description",
            value=st.session_state["saved_jd"],
            height=420,
            label_visibility="collapsed",
            placeholder="Paste the full JD here…",
        )

    if st.button("✨ Generate Tailored Resume", type="primary", use_container_width=True):
        if not api_key:
            st.error("Need a Google API key to generate.")
        elif not jd.strip():
            st.warning("Please paste a job description.")
        elif not base.strip():
            st.warning("Base resume is empty. Click 'Restore default base resume' to bring it back.")
        else:
            st.session_state["saved_jd"] = jd
            st.session_state["saved_base"] = base
            with st.spinner("Tailoring resume to JD…"):
                model_out = call_gemini(api_key, jd, resume_text=base)
                if "error" in model_out:
                    st.error(model_out["error"])
                else:
                    final = assemble_resume(model_out)
                    st.session_state["tailored"] = final
                    st.rerun()

else:
    data = st.session_state["tailored"]

    # Top bar
    c1, c2 = st.columns([4, 1])
    with c1:
        st.markdown(f"## 🎯 Target: {data['target_company']}")
        st.caption(f"Tailored title: **{data['candidate_title']}**")
        archetype = data.get("jd_archetype", "")
        if archetype:
            archetype_display = {
                "data_analyst": "🟢 Data Analyst (strong fit)",
                "bi_analyst": "🟢 BI Analyst (strong fit)",
                "junior_data_scientist": "🟡 Junior Data Scientist (moderate fit, check warning)",
                "analytics_engineer": "🟡 Analytics Engineer (moderate fit, check warning)",
                "graduate_programme_tech": "🟢 Graduate Programme, Tech/Data (strong fit)",
                "graduate_programme_business": "🟢 Graduate Programme, Business/Consulting (strong fit)",
                "qa_automation": "🟡 QA Automation (weak fit, see warning)",
                "unknown_stretch": "🔴 Stretch role, see warning",
            }.get(archetype, archetype)
            st.caption(f"JD archetype: {archetype_display}")
    with c2:
        if st.button("New JD", use_container_width=True):
            st.session_state["tailored"] = None
            st.session_state["saved_jd"] = ""
            st.session_state["cover_letter"] = None
            st.rerun()

    # JD fit warning
    fit_warning = data.get("jd_fit_warning", "").strip()
    if fit_warning:
        st.warning(f"⚠️ **JD fit warning:** {fit_warning}")

    # Tabs
    tab_preview, tab_edit, tab_cover, tab_download = st.tabs(
        ["👀 Preview", "📝 Edit", "✍️ Cover Letter", "📥 Download"]
    )

    with tab_preview:
        st.subheader("Professional Profile")
        st.write(data["summary"])

        st.subheader("Key Skills, Tools and Technologies")
        for cat, skill_list in data["skills"].items():
            st.markdown(f"- **{cat}:** {', '.join(skill_list)}")

        st.subheader("Professional Experience")
        for role in data["experience"]:
            st.markdown(
                f"**{role['role_title']}** | {role['company']} | "
                f"{role['location']} | {role['dates']}"
            )
            for r in role["responsibilities"]:
                st.markdown(f"- {r}")
            if role["achievements"]:
                st.markdown("**Achievements:**")
                for a in role["achievements"]:
                    st.markdown(f"- {a}")

        st.subheader("Projects")
        for proj in data["projects"]:
            st.markdown(
                f"**{proj['title']}** | {proj['context']} | {proj['dates']}"
            )
            for b in proj["bullets"]:
                st.markdown(f"- {b}")

        st.subheader("Education")
        for edu in data["education"]:
            st.markdown(f"**{edu['degree']}**")
            st.markdown(f"{edu['institution']} | {edu['dates']} | {edu['grade']}")
            if edu.get("extra"):
                st.write(edu["extra"])

        st.subheader("Certifications")
        for c in data["certifications"]:
            st.markdown(f"- {c}")

        st.subheader("Additional Information")
        for a in data["additional_info"]:
            st.markdown(f"- {a}")

    with tab_edit:
        with st.form("edit_form"):
            data["candidate_title"] = st.text_input("Title under name", data["candidate_title"])
            data["summary"] = st.text_area("Summary", data["summary"], height=160)

            st.markdown("##### Skills (comma-separated per category)")
            for cat in list(data["skills"].keys()):
                joined = ", ".join(data["skills"][cat])
                edited = st.text_area(cat, joined, height=70, key=f"sk_{cat}")
                data["skills"][cat] = [s.strip() for s in edited.split(",") if s.strip()]

            st.markdown("##### Experience")
            for i, role in enumerate(data["experience"]):
                with st.expander(f"{role['role_title']} @ {role['company']}", expanded=False):
                    resps_text = "\n".join(role["responsibilities"])
                    new_resps = st.text_area("Responsibilities (one per line)",
                                             resps_text, height=180, key=f"r_{i}")
                    role["responsibilities"] = [
                        line.strip() for line in new_resps.split("\n") if line.strip()
                    ]
                    achs_text = "\n".join(role["achievements"])
                    new_achs = st.text_area("Achievements (one per line)",
                                            achs_text, height=80, key=f"a_{i}")
                    role["achievements"] = [
                        line.strip() for line in new_achs.split("\n") if line.strip()
                    ]

            st.markdown("##### Projects")
            for i, proj in enumerate(data["projects"]):
                with st.expander(f"{proj['title']}", expanded=False):
                    bullets_text = "\n".join(proj["bullets"])
                    new_bullets = st.text_area("Bullets (one per line)",
                                               bullets_text, height=120, key=f"p_{i}")
                    proj["bullets"] = [
                        line.strip() for line in new_bullets.split("\n") if line.strip()
                    ]

            if st.form_submit_button("💾 Save edits", type="primary"):
                st.session_state["tailored"] = data
                st.success("Saved.")
                st.rerun()

    with tab_cover:
        st.caption(
            "A short, human-sounding cover letter. Generated separately so it "
            "doesn't slow down resume tailoring."
        )

        cover_btn_label = "✨ Draft Cover Letter" if not st.session_state["cover_letter"] else "🔄 Re-draft Cover Letter"
        if st.button(cover_btn_label, type="primary"):
            if not api_key:
                st.error("Need a Google API key to generate.")
            elif not st.session_state["saved_jd"].strip():
                st.warning("No saved JD found. Generate the resume first.")
            else:
                with st.spinner("Drafting cover letter…"):
                    cl = generate_cover_letter(api_key, data, st.session_state["saved_jd"])
                    if cl.startswith("ERROR:"):
                        st.error(cl)
                    else:
                        st.session_state["cover_letter"] = cl
                        st.rerun()

        if st.session_state["cover_letter"]:
            edited = st.text_area(
                "Cover letter (editable)",
                st.session_state["cover_letter"],
                height=420,
            )
            st.session_state["cover_letter"] = edited

            try:
                cl_bytes = render_cover_letter_docx(
                    st.session_state["cover_letter"],
                    target_company=data.get("target_company", ""),
                )
                company_safe = re.sub(r"[^A-Za-z0-9_-]", "_",
                                      (data["target_company"] or "Company").strip()) or "Company"
                st.download_button(
                    "📄 Download Cover Letter (.docx)",
                    data=cl_bytes,
                    file_name=f"CoverLetter_Vishnunath_Nharekkat_{company_safe}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                )
            except Exception as e:
                st.error(f"Cover letter render error: {e}")

    with tab_download:
        company = data["target_company"] or "Company"
        company_safe = re.sub(r"[^A-Za-z0-9_-]", "_", company.strip()) or "Company"
        filename_base = f"Vishnunath_Nharekkat_{company_safe}"

        st.text_input("Filename (no extension)", filename_base, key="fname")
        fname = st.session_state.get("fname", filename_base)

        c1, c2 = st.columns(2)
        try:
            docx_bytes = render_docx(data)
            c1.download_button(
                "📄 Word (.docx)",
                data=docx_bytes,
                file_name=f"{fname}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )
        except Exception as e:
            c1.error(f"DOCX error: {e}")

        try:
            pdf_bytes = render_pdf(data)
            c2.download_button(
                "📕 PDF",
                data=pdf_bytes,
                file_name=f"{fname}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            c2.error(f"PDF error: {e}")
